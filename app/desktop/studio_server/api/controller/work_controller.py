import glob
import json
import os
import re
from typing import List
from pathlib import Path
from docx import Document
from docx.document import Document as DocumentObject
from fastapi import FastAPI, File, UploadFile
from fastapi.responses import JSONResponse, StreamingResponse
from ollama import chat as ollama_chat
from openai import OpenAI
from app.desktop.studio_server.api.service.work_service import WorkService
from app.desktop.studio_server.datamodel.dto.LLMdto import ChatRequest
from app.desktop.studio_server.datamodel.dto.Workdto import WorkRunParams
from app.desktop.studio_server.utils.config import Config
from app.desktop.studio_server.api.controller.model_controller import llmProviderManager
from app.desktop.studio_server.utils.exhaustive_error import error_code_to_message
def default_share_path():
    return os.path.join(Path.home(), ".XHAgentData","work")

content = ""
keywords =[]
async def llm_chat(model: str,provider: str,message:str,prompt: str=""):
    stream = None
    if provider == "ollama":
        stream = ollama_chat(
            model=model,
            options={
                'temperature': 0.4,      # 生成的随机性，数值越高越随机
                'top_k': 40,             # 采样时只考虑概率最高的前 k 个 token
                'top_p': 0.9,            # Nucleus 采样，控制高概率 token 的总占比
                'repeat_penalty': 1.1,   # **重复惩罚**，数值越高，重复内容概率越低
            },
            messages=[
                # {"system":"assistant","content":prompt},
                {"role": "user", "content": prompt+message}
            ],
            stream=True,
        ) 
        for chunk in stream:
            yield chunk['message']['content']
    else:
        if provider == "deepseek":
            match model:
                case "DeepSeek-V3":
                    model = "deepseek-chat"
                case "DeepSeek-R1":
                    model = "deepseek-reasoner"
        client: OpenAI = llmProviderManager.get_clients()[provider].clients
        stream = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "user", "content": prompt + message}
            ],
            stream=True
        )
        for chunk in stream:
            content = chunk.choices[0].delta.content  # 获取内容
            if content:  # 只 yield 非空内容
                yield content    
        
    
def check_docx(docx: DocumentObject):
    global content, keywords
    check_content = ""
    """检查文档中是否包含所有关键字"""
    for para in docx.paragraphs:
        check_content += para.text
    matches = re.findall(r'【[^】]*】', check_content)
    if len(matches) != 6:
        return False
    else:
        keywords = matches
    capture = False
    start_keyword = matches[0]
    end_keyword = matches[3]
    # 获取文档中所有段落的文本内容
    for para in docx.paragraphs:
        if start_keyword in para.text:
            capture = True
            content += para.text + "\n"
        elif capture:
            content += para.text + "\n"
            if end_keyword in para.text:
                break
    return True

def connect_work_api(app: FastAPI):
    if not os.path.exists(default_share_path()):
        os.makedirs(default_share_path())
    @app.post("/api/work/upload/docx")
    def update_docx_check(file: UploadFile = File(...)):
        doc = Document(file.file)  # 直接传递 file 流
        if check_docx(doc):
            doc.save(os.path.join(default_share_path(), "temp.docx"))
            return True
        else:
            return False
    
    @app.post("/api/work/upload/image")
    async def update_docx(files: List[UploadFile] = File(...)):
        saved_files = []
        for file in files:
            file_path = os.path.join(default_share_path(), file.filename)
            os.makedirs(os.path.dirname(file_path), exist_ok=True)  # 确保目录存在
            # 如果图片存在就跳过
            if os.path.exists(file_path):
                continue
            with open(file_path, "wb") as f:
                # 逐块读取文件并写入
                while chunk := await file.read(1024):
                    f.write(chunk)
            
                saved_files.append({
                    "filename": file.filename,
                    "msg": f"{file.filename} 上传成功",
                })

        return {"uploaded_files": saved_files}

    @app.post("/api/work/summary")
    async def gen_summary(chatParams: ChatRequest):
        try:
            model = chatParams.model
            provider = chatParams.provider
            if content == "":
                return StreamingResponse("文档内无描述任务，请重写上传", media_type="text/plain")
            prompt = "请根据以下内容写有一份150字以内的实验报告心得（不要使用“首先”、“其次”、“然而”、“总的来说”、“总之”这些副词）。尽量用主动句，增加文章力量。避免使用陈词滥调，换成新颖的表达。避免生硬的总结或说教。删除小标题。"

            async def safe_generator():
                try:
                    async for chunk in llm_chat(model, provider, content, prompt):
                        yield chunk
                except Exception as e:
                    error_msg = str(e)
                    error_code = None
                    
                    # 尝试解析错误码
                    if "Error code:" in error_msg:
                        try:
                            error_code = int(error_msg.split("Error code:")[1].split("-")[0].strip())
                        except ValueError:
                            pass  # 如果解析失败，就不处理
                        
                    # 根据错误码返回对应的错误描述
                    if error_code:
                        yield f"{error_code_to_message(error_code)}"
                    else:
                        yield f"【未知错误】{error_msg}"

            return StreamingResponse(safe_generator(), media_type="text/plain")

        except Exception as e:
            return JSONResponse(content={"error": str(e)}, status_code=500)

    
    @app.post("/api/work/rewrite")
    async def rewrite(chatParams: ChatRequest):
        try:
            model = chatParams.model
            provider = chatParams.provider
            prompt = "请帮我把下面的内容进行重写。注意保持原句的格式。"
            if chatParams.message == "":
                return StreamingResponse("无描述", media_type="text/plain")
            async def safe_generator():
                    try:
                        async for chunk in llm_chat(model, provider, chatParams.message, prompt):
                            yield chunk
                    except Exception as e:
                        error_msg = str(e)
                        error_code = None
                        
                        # 尝试解析错误码
                        if "Error code:" in error_msg:
                            try:
                                error_code = int(error_msg.split("Error code:")[1].split("-")[0].strip())
                            except ValueError:
                                pass  # 如果解析失败，就不处理
                            
                        # 根据错误码返回对应的错误描述
                        if error_code:
                            yield f"{error_code_to_message(error_code)}"
                        else:
                            yield f"【未知错误】{error_msg}"

            return StreamingResponse(safe_generator(), media_type="text/plain")
        except Exception as e:
            return JSONResponse(content={"error": str(e)}, status_code=500)
    
    @app.get("/api/work/delete/image")
    async def delete_image(image_name: str):
        image_path = os.path.join(default_share_path(), image_name)
        if os.path.exists(image_path):
            os.remove(image_path)

    @app.get("/api/work/setting")
    async def get_work_setting():
        return {
            "name": Config.shared().get("name"),
            "course": Config.shared().get("course"),
            "id": Config.shared().get("id")
        }
    
    workService = WorkService()
    @app.post("/api/work/run")
    async def run_work(workRunParams: WorkRunParams):
        global keywords
        return await workService.run(workRunParams,keywords)
    
    def format_size(size_in_bytes):
        units = ["B", "KB", "MB", "GB", "TB"]
        size = size_in_bytes
        unit_index = 0

        while size >= 1024 and unit_index < len(units) - 1:
            size /= 1024
            unit_index += 1

        return f"{size:.2f} {units[unit_index]}"
    
    @app.get("/api/work/cache")
    def read_cache_info():
        try:
            total_size = 0
            image_files = []
            extensions = ("jpg", "png", "jpeg", "gif", "bmp", "tiff", "webp")

            for ext in extensions:
                image_files.extend(glob.glob(os.path.join(workService.default_share_path(), f"*.{ext}")))

            total_size = sum(os.path.getsize(img) for img in image_files)

            return JSONResponse(content={
                'fileCount': len(image_files),
                'total_size': format_size(total_size)
            })
        except Exception as e:
            return JSONResponse(content={'fileCount': 0, 'total_size': "0 B"}, status_code=500)

        
        
    @app.get("/api/work/deleteCache")
    def deleteCache():
        extensions=("jpg", "png", "jpeg", "gif", "bmp", "tiff", "webp")
        try:
            for ext in extensions:
                for img_file in glob.glob(os.path.join(workService.default_share_path(), f"*.{ext}")):
                    os.remove(img_file)
            return True
        except Exception as e:
            return False
    