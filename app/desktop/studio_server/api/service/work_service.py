import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.document import Document as DocumentObject
from fastapi.responses import FileResponse

from app.desktop.studio_server.datamodel.dto.Workdto import WorkRunParams
from app.desktop.studio_server.utils.config import Config


class WorkService:
    def __init__(self):
        self.doc  = None
    
    def default_share_path(self):
        return os.path.join(Path.home(), ".XHAgentData","work")
    
    def default_work_doc_path(self):
        return os.path.join(self.default_share_path(), "temp.docx")
    
    def initDocx(self):
        self.doc = Document(self.default_work_doc_path())
        
    # 修改页眉
    def modify_header(self,name,_id):
        new_header = f"实验人：{name}（学号：{_id}）"
        # 修改每个节的页眉
        for section in self.doc.sections:
            header = section.header
            # 清空原有内容
            header.paragraphs[0].clear()
            # 添加新的内容
            run =header.paragraphs[0].add_run(new_header)
            run.font.size = Pt(10.5)  # 5号字（10.5磅）
            run.font.name = "宋体"
            
            
    # 插入心得
    def insert_learning(self,summary:str,keywords:list[str]):
        # 添加心得栏
        text_paragraph = self.doc.add_paragraph()
        text_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 水平居左
        # 设置段落间距为 1.5 倍行距
        text_paragraph.paragraph_format.line_spacing = Pt(
            24
        )  # 1.5倍行距（16磅*1.5=24磅）
        run = text_paragraph.add_run(keywords[-1])
        run.font.size = Pt(10.5)  # 5号字（10.5磅）
        run.font.name = "宋体"  # 设置字体为宋体
        run.font.bold = True

        def format_text(content:str):
            paragraphs = content.split("\n\n")  # 将内容按段落分开
            for paragraph_text in paragraphs:
                text_paragraph = self.doc.add_paragraph()  # 使用文档对象添加文字段落
                text_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 水平居左

                # 设置段落间距为 1.5 倍行距
                text_paragraph.paragraph_format.line_spacing = Pt(
                    24
                ) 

                # 设置段落开头缩进2个字（24磅，约等于2个汉字宽度）
                text_paragraph.paragraph_format.first_line_indent = Pt(24)

                # 添加段落内容
                run = text_paragraph.add_run(paragraph_text)
                run.font.size = Pt(10.5)  # 设置字体为5号字（10.5磅）
                run.font.name = "宋体"  # 设置字体为宋体

        format_text(summary)

    # 文档清洗
    def doc_clean(self,keywords:list[str]):
        # 标记是否找到类似 "【实验过程记录】"
        found = False
        # 遍历文档中的段落
        for para in self.doc.paragraphs:
            if found:
                CT_P = para._element
                CT_P.getparent().remove(CT_P)
            elif keywords[4] in para.text:
                found = True
                continue
   
    # 插入实验过程
    def insert_paragraphs(self,pic_num, imagePath, text):
        n_pic_num= 0
        if text !="":
            text_paragraph = self.doc.add_paragraph()
            run = text_paragraph.add_run(text)  
            run.font.size = Pt(10.5)  # 5号字（10.5磅）
            run.font.name = "宋体" 
            text_paragraph.paragraph_format.left_indent = Inches(0.5) 

            if imagePath !="":  
                suffix = f"如图 {pic_num}所示" if text[-1] in "。，.," else f"，如图 {pic_num}所示"
                run = text_paragraph.add_run(suffix)
                run.font.size = Pt(10.5)  
                run.font.name = "宋体"

            # 设置段落格式
            text_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  
            text_paragraph.paragraph_format.line_spacing = Pt(24)  
            self.doc.add_paragraph().paragraph_format.space_after = Pt(0.2)
        if imagePath != "": # 如果有图片
            # 插入图片
            image_paragraph = self.doc.add_paragraph()  # 使用文档对象添加新段落
            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_paragraph.add_run().add_picture(imagePath, width=Inches(5))  # 调整宽度
            # 添加图片题注（段落）
            self.doc.add_paragraph()
            caption_paragraph = self.doc.add_paragraph()
            run = caption_paragraph.add_run(f"图 {pic_num}")
            run.font.size = Pt(10.5)  # 5号字（10.5磅
            caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            n_pic_num +=1
        return n_pic_num
            
    async def run(self,workRunParams:WorkRunParams,keywords:list[str]):
        name = workRunParams.name
        course = workRunParams.course
        _id = workRunParams.id
        
        self.initDocx()
         # 修改页眉
        self.modify_header(name,_id)
        # 文档清洗
        self.doc_clean(keywords)
        
        pic_num = 1
        for image_group in workRunParams.images_texts:
            if image_group.imageName != "":
                imagePath = os.path.join(self.default_share_path(),image_group.imageName)
            else:
                imagePath = ""
            # 插入实验过程
            pic_num+=self.insert_paragraphs(pic_num,imagePath, image_group.text)
        # 插入心得
        self.insert_learning(workRunParams.summary,keywords)
        
        Config.shared().save_setting("name",name)
        Config.shared().save_setting("course",course)
        Config.shared().save_setting("id",_id)
        
        self.doc.save(self.default_work_doc_path())
        return FileResponse(
            path=self.default_work_doc_path(),
            filename=f"{course}_{name}_{_id}_{workRunParams.docxName}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )